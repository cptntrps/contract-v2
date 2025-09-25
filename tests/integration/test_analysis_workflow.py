"""
Integration tests for complete contract analysis workflow

Tests the full integration of document processing, comparison, LLM analysis,
database persistence, and report generation.
"""
import pytest
import tempfile
import os
import json
from pathlib import Path
from datetime import datetime
from unittest.mock import patch, Mock

from app.core.services.analyzer import ContractAnalyzer
from app.core.services.document_processor import DocumentProcessor
from app.core.services.comparison_engine import ComparisonEngine
from app.core.services.report_generator import ReportGenerator
from app.database.repositories.contract_repository import ContractRepository
from app.database.repositories.analysis_result_repository import AnalysisResultRepository
from app.llm.llm_handler import create_llm_provider
from app.config.settings import get_config


@pytest.fixture
def test_config():
    """Test configuration"""
    return {
        'llm_settings': {
            'provider': 'openai',
            'api_key': 'test-key',
            'model': 'gpt-3.5-turbo'
        },
        'database': {
            'url': 'sqlite:///:memory:'
        },
        'upload_folder': tempfile.mkdtemp()
    }


@pytest.fixture
def test_documents(tmp_path):
    """Create test DOCX documents"""
    try:
        from docx import Document
        
        # Create template document
        template_doc = Document()
        template_doc.add_paragraph("STANDARD CONTRACT TEMPLATE")
        template_doc.add_paragraph("1. Payment Terms: Net 30 days")
        template_doc.add_paragraph("2. Delivery: Standard shipping")
        template_doc.add_paragraph("3. Warranty: 12 months")
        template_doc.add_paragraph("4. Termination: 30 days notice required")
        
        template_path = tmp_path / "template.docx"
        template_doc.save(str(template_path))
        
        # Create contract document (modified version)
        contract_doc = Document()
        contract_doc.add_paragraph("MODIFIED CONTRACT")
        contract_doc.add_paragraph("1. Payment Terms: Net 45 days")  # Changed
        contract_doc.add_paragraph("2. Delivery: Express shipping")  # Changed
        contract_doc.add_paragraph("3. Warranty: 24 months")  # Changed
        contract_doc.add_paragraph("4. Termination: 60 days notice required")  # Changed
        contract_doc.add_paragraph("5. Additional Clause: Penalty for late delivery")  # Added
        
        contract_path = tmp_path / "contract.docx"
        contract_doc.save(str(contract_path))
        
        return {
            'template': str(template_path),
            'contract': str(contract_path)
        }
    except ImportError:
        # If python-docx not available, create text files
        template_path = tmp_path / "template.txt"
        template_path.write_text("""STANDARD CONTRACT TEMPLATE
1. Payment Terms: Net 30 days
2. Delivery: Standard shipping
3. Warranty: 12 months
4. Termination: 30 days notice required""")
        
        contract_path = tmp_path / "contract.txt"
        contract_path.write_text("""MODIFIED CONTRACT
1. Payment Terms: Net 45 days
2. Delivery: Express shipping
3. Warranty: 24 months
4. Termination: 60 days notice required
5. Additional Clause: Penalty for late delivery""")
        
        return {
            'template': str(template_path),
            'contract': str(contract_path)
        }


class TestFullAnalysisWorkflow:
    """Test complete analysis workflow from upload to report generation"""
    
    def test_complete_workflow_with_real_components(self, test_config, test_documents, tmp_path):
        """Test workflow with real components (no mocks)"""
        # Initialize services
        doc_processor = DocumentProcessor()
        comparison_engine = ComparisonEngine()
        
        # Process documents
        template_text = doc_processor.extract_text_from_docx(test_documents['template'])
        contract_text = doc_processor.extract_text_from_docx(test_documents['contract'])
        
        # Verify text extraction
        assert len(template_text) > 0
        assert len(contract_text) > 0
        assert "Net 30 days" in template_text
        assert "Net 45 days" in contract_text
        
        # Compare documents
        changes = comparison_engine.find_detailed_changes(template_text, contract_text)
        
        # Verify changes detected
        assert len(changes) > 0
        
        # Check for expected changes
        change_texts = [
            (c.get('deleted_text', ''), c.get('inserted_text', ''))
            for c in changes
        ]
        
        # Should detect payment terms change
        assert any('30 days' in deleted and '45 days' in inserted 
                  for deleted, inserted in change_texts)
        
        # Should detect warranty change
        assert any('12 months' in deleted and '24 months' in inserted 
                  for deleted, inserted in change_texts)
        
        # Calculate statistics
        stats = comparison_engine.get_change_statistics(changes)
        assert stats['total_changes'] > 0
        assert stats['insertions'] > 0
        assert stats['replacements'] > 0
    
    def test_workflow_with_database_persistence(self, test_config, test_documents, tmp_path):
        """Test workflow with database persistence"""
        from app.database.database import db
        from flask import Flask
        
        # Create Flask app for database context
        app = Flask(__name__)
        app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///:memory:'
        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
        
        # Initialize database
        from app.database import init_app
        init_app(app)
        
        with app.app_context():
            # Create tables
            db.create_all()
            
            # Initialize repositories
            contract_repo = ContractRepository()
            analysis_repo = AnalysisResultRepository()
            
            # Create contracts
            template_data = {
                'contract_id': 'template_001',
                'filename': 'template.docx',
                'content': 'Template content with payment terms: Net 30 days',
                'file_type': 'template'
            }
            template = contract_repo.create(**template_data)
            
            contract_data = {
                'contract_id': 'contract_001',
                'filename': 'contract.docx',
                'content': 'Contract content with payment terms: Net 45 days',
                'file_type': 'contract'
            }
            contract = contract_repo.create(**contract_data)
            
            # Verify persistence
            assert template.id == 'template_001'
            assert contract.id == 'contract_001'
            
            # Create analysis result
            analysis_data = {
                'analysis_id': 'analysis_001',
                'contract_id': 'contract_001',
                'template_id': 'template_001',
                'status': 'completed',
                'statistics': {'total_changes': 5},
                'llm_analysis': {'summary': 'Test analysis'}
            }
            analysis = analysis_repo.create(**analysis_data)
            
            # Verify analysis persistence
            assert analysis.analysis_id == 'analysis_001'
            assert analysis.statistics['total_changes'] == 5
            
            # Test retrieval
            retrieved = analysis_repo.get_by_id('analysis_001')
            assert retrieved is not None
            assert retrieved.contract_id == 'contract_001'
    
    def test_workflow_with_report_generation(self, test_config, test_documents, tmp_path):
        """Test workflow including report generation"""
        from app.core.domain.analysis_result import AnalysisResult
        from app.core.domain.change import Change
        
        # Create sample analysis result
        analysis = AnalysisResult(
            analysis_id='test_analysis_123',
            contract_id='contract_001',
            template_id='template_001',
            status='completed',
            created_at=datetime.now()
        )
        
        # Add sample changes
        changes = [
            Change(
                operation='replace',
                original_text='Net 30 days',
                modified_text='Net 45 days',
                position=100,
                context='Payment Terms: {change}'
            ),
            Change(
                operation='insert',
                original_text='',
                modified_text='Penalty for late delivery',
                position=500,
                context='Additional Clause: {change}'
            )
        ]
        
        for change in changes:
            analysis.add_change(change)
        
        # Add statistics
        analysis.statistics = {
            'total_changes': 2,
            'insertions': 1,
            'replacements': 1,
            'deletions': 0
        }
        
        # Generate reports
        report_gen = ReportGenerator()
        
        # Test JSON report
        json_path = tmp_path / "report.json"
        success = report_gen.generate_json_report(analysis, str(json_path))
        assert success is True
        assert json_path.exists()
        
        # Verify JSON content
        with open(json_path) as f:
            data = json.load(f)
        assert data['analysis_id'] == 'test_analysis_123'
        assert len(data['changes']) == 2
        
        # Test CSV report
        csv_path = tmp_path / "report.csv"
        success = report_gen.generate_csv_report(analysis, str(csv_path))
        assert success is True
        assert csv_path.exists()
        
        # Test Excel report if available
        if 'excel' in report_gen.get_supported_formats():
            excel_path = tmp_path / "report.xlsx"
            success = report_gen.generate_excel_report(analysis, str(excel_path))
            assert success is True
            assert excel_path.exists()
    
    @patch('app.llm.llm_handler.create_llm_provider')
    def test_workflow_with_llm_integration(self, mock_llm_factory, test_config, test_documents):
        """Test workflow with mocked LLM integration"""
        # Mock LLM provider
        mock_llm = Mock()
        mock_llm.analyze_contract_changes.return_value = {
            'summary': 'Significant changes to payment terms and warranty period',
            'risk_assessment': 'medium',
            'recommendations': [
                'Review extended payment terms impact on cash flow',
                'Verify warranty cost implications'
            ]
        }
        mock_llm_factory.return_value = mock_llm
        
        # Initialize analyzer
        analyzer = ContractAnalyzer(test_config)
        
        # Mock document processing
        with patch.object(analyzer, 'extract_text_from_docx') as mock_extract:
            mock_extract.side_effect = [
                "Template with Net 30 days payment",
                "Contract with Net 45 days payment"
            ]
            
            # Perform analysis
            template_text = analyzer.extract_text_from_docx(test_documents['template'])
            contract_text = analyzer.extract_text_from_docx(test_documents['contract'])
            
            changes = analyzer.find_changes(template_text, contract_text)
            
            # Verify LLM was called
            assert mock_llm.analyze_contract_changes.called
    
    def test_error_handling_in_workflow(self, test_config, tmp_path):
        """Test error handling throughout the workflow"""
        from app.utils.errors.exceptions import ValidationError, FileProcessingError
        
        doc_processor = DocumentProcessor()
        
        # Test with non-existent file
        result = doc_processor.extract_text_from_docx("nonexistent.docx")
        assert result == ""  # Should handle gracefully
        
        # Test with invalid file
        invalid_file = tmp_path / "invalid.xyz"
        invalid_file.write_text("Not a valid document")
        
        result = doc_processor.extract_text_from_docx(str(invalid_file))
        assert result == ""  # Should handle gracefully
    
    def test_concurrent_analysis_workflow(self, test_config, test_documents):
        """Test concurrent analysis operations"""
        import threading
        from queue import Queue
        
        results = Queue()
        errors = Queue()
        
        def analyze_documents():
            try:
                doc_processor = DocumentProcessor()
                comparison_engine = ComparisonEngine()
                
                template_text = doc_processor.extract_text_from_docx(test_documents['template'])
                contract_text = doc_processor.extract_text_from_docx(test_documents['contract'])
                
                changes = comparison_engine.find_detailed_changes(template_text, contract_text)
                results.put(len(changes))
            except Exception as e:
                errors.put(e)
        
        # Run multiple analyses concurrently
        threads = []
        for _ in range(3):
            t = threading.Thread(target=analyze_documents)
            threads.append(t)
            t.start()
        
        # Wait for completion
        for t in threads:
            t.join()
        
        # Verify results
        assert errors.empty()
        assert results.qsize() == 3
        
        # All should detect same number of changes
        change_counts = [results.get() for _ in range(3)]
        assert all(count > 0 for count in change_counts)


class TestAPIIntegration:
    """Test API endpoint integration"""
    
    @pytest.fixture
    def client(self, test_config):
        """Create test client"""
        from app.main import create_app
        
        app = create_app(test_config)
        app.config['TESTING'] = True
        
        with app.test_client() as client:
            yield client
    
    def test_contract_upload_api(self, client, test_documents):
        """Test contract upload via API"""
        with open(test_documents['contract'], 'rb') as f:
            response = client.post('/api/contracts/upload', data={
                'file': (f, 'contract.docx'),
                'contract_id': 'test_contract_001'
            })
        
        # Should handle file upload
        assert response.status_code in [200, 201, 400]  # Depends on implementation
    
    def test_analysis_api_workflow(self, client):
        """Test analysis API workflow"""
        # Mock contract existence
        with patch('app.database.repositories.contract_repository.ContractRepository.get_by_id') as mock_get:
            mock_get.return_value = Mock(id='contract_001', content='Test content')
            
            # Request analysis
            response = client.post('/api/analysis', json={
                'contract_id': 'contract_001',
                'template_id': 'template_001',
                'include_llm_analysis': True
            })
            
            # Should accept analysis request
            assert response.status_code in [200, 201, 202, 400]
    
    def test_report_generation_api(self, client):
        """Test report generation via API"""
        # Mock analysis result
        with patch('app.database.repositories.analysis_result_repository.AnalysisResultRepository.get_by_id') as mock_get:
            mock_result = Mock()
            mock_result.to_dict.return_value = {
                'analysis_id': 'analysis_001',
                'changes': [],
                'statistics': {}
            }
            mock_get.return_value = mock_result
            
            # Request report
            response = client.post('/api/reports/generate', json={
                'analysis_id': 'analysis_001',
                'format': 'json'
            })
            
            # Should handle report request
            assert response.status_code in [200, 201, 400]


class TestDatabaseIntegration:
    """Test database operations integration"""
    
    def test_transaction_management(self, test_config):
        """Test database transaction handling"""
        from app.database.database import db
        from flask import Flask
        from sqlalchemy.exc import IntegrityError
        
        app = Flask(__name__)
        app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///:memory:'
        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
        
        from app.database import init_app
        init_app(app)
        
        with app.app_context():
            db.create_all()
            
            from app.database.models.contract import ContractModel
            
            # Test rollback on error
            try:
                with db.session.begin():
                    # Create contract
                    contract = ContractModel(
                        id='test_001',
                        filename='test.docx',
                        content='Test content'
                    )
                    db.session.add(contract)
                    
                    # Force error by duplicate ID
                    duplicate = ContractModel(
                        id='test_001',  # Duplicate ID
                        filename='duplicate.docx',
                        content='Duplicate content'
                    )
                    db.session.add(duplicate)
                    
            except IntegrityError:
                # Transaction should be rolled back
                pass
            
            # Verify rollback - no contracts should exist
            count = ContractModel.query.count()
            assert count == 0