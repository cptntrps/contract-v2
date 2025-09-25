"""
Word Report Generator - Domain Layer

Specialized generator for Word documents with optional track changes.
Following architectural standards: single responsibility, comprehensive documentation.
"""

import os
from typing import Dict, Any, Optional

from .base_generator import BaseReportGenerator, ReportError
from ....utils.logging.setup import get_logger

# Word handling with graceful fallback
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Windows COM for track changes (optional)
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

logger = get_logger(__name__)


class WordReportGenerator(BaseReportGenerator):
    """
    Specialized Word document generator with optional track changes.
    
    Purpose: Generates Word documents with professional formatting and
    optional track changes functionality using COM automation on Windows.
    Handles Word-specific features like tables, styles, and revision tracking.
    
    AI Context: Word-specific report generator. Uses python-docx for basic
    documents and win32com for track changes on Windows. When debugging,
    check dependency availability and COM automation status.
    """
    
    @property
    def format_name(self) -> str:
        """Return format name for Word documents."""
        return 'word'
    
    @property
    def file_extension(self) -> str:
        """Return file extension for Word documents."""
        return '.docx'
    
    @property
    def dependencies_available(self) -> bool:
        """Check if python-docx is available."""
        return DOCX_AVAILABLE
    
    def generate_report(self, analysis: Any, output_path: str,
                       enable_track_changes: bool = False,
                       include_summary: bool = True,
                       **options) -> bool:
        """
        Generate Word report with optional track changes.
        
        Purpose: Creates professional Word document with analysis data,
        formatted tables, and optional track changes using COM automation.
        Falls back to python-docx if COM is unavailable.
        
        Args:
            analysis: Analysis result containing changes and metadata
            output_path: Path to save Word document
            enable_track_changes: Whether to enable track changes via COM
            include_summary: Whether to include summary section
            **options: Additional Word-specific options
            
        Returns:
            True if successful
            
        Raises:
            ReportError: If Word generation fails or dependencies missing
            
        AI Context: Main Word document generation method. Tries COM automation
        first for track changes, falls back to python-docx. For debugging,
        check COM availability and document structure creation.
        """
        # Validate dependencies and path
        self.validate_dependencies()
        normalized_path = self.validate_output_path(output_path)
        
        self.log_generation_start(normalized_path)
        
        try:
            # Try COM automation if requested and available
            if enable_track_changes and WIN32COM_AVAILABLE:
                success = self._generate_with_com_automation(analysis, normalized_path, **options)
                if success:
                    self.log_generation_success(normalized_path)
                    return True
                else:
                    logger.warning("COM automation failed, falling back to python-docx")
            
            # Use python-docx for standard Word document
            return self._generate_with_python_docx(analysis, normalized_path, include_summary, **options)
            
        except Exception as e:
            self.log_generation_failure(normalized_path, e)
            raise ReportError(f"Failed to generate Word report: {e}")
    
    def _generate_with_python_docx(self, analysis: Any, output_path: str,
                                  include_summary: bool, **options) -> bool:
        """
        Generate Word document using python-docx library.
        
        Args:
            analysis: Analysis result data
            output_path: Path to save document
            include_summary: Whether to include summary section
            **options: Additional options
            
        Returns:
            True if successful
        """
        doc = Document()
        
        # Document title
        title = doc.add_heading('Contract Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document metadata
        self._add_metadata_section(doc, analysis)
        
        # Add summary section if requested
        if include_summary:
            self._add_summary_section(doc, analysis)
        
        # Add statistics section
        self._add_statistics_section(doc, analysis)
        
        # Add detailed changes section
        self._add_changes_table(doc, analysis)
        
        # Save document
        doc.save(output_path)
        self.log_generation_success(output_path)
        return True
    
    def _generate_with_com_automation(self, analysis: Any, output_path: str,
                                    **options) -> bool:
        """
        Generate Word document with track changes using COM automation.
        
        Args:
            analysis: Analysis result data
            output_path: Path to save document
            **options: Additional options
            
        Returns:
            True if successful, False if COM automation fails
            
        AI Context: Uses Windows COM to create Word document with proper
        track changes. This provides true revision tracking but only works
        on Windows with Office installed.
        """
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Create new document
            doc = word.Documents.Add()
            doc.TrackRevisions = True
            
            # Get selection object for content manipulation
            selection = word.Selection
            
            # Add title
            selection.Font.Size = 18
            selection.Font.Bold = True
            selection.TypeText("Contract Analysis Report\n\n")
            selection.Font.Size = 12
            selection.Font.Bold = False
            
            # Add metadata
            metadata_items = [
                f"Analysis ID: {getattr(analysis, 'analysis_id', 'N/A')}",
                f"Contract ID: {getattr(analysis, 'contract_id', 'N/A')}",
                f"Template ID: {getattr(analysis, 'template_id', 'N/A')}",
                f"Status: {getattr(analysis, 'status', 'N/A')}\n\n"
            ]
            
            for item in metadata_items:
                selection.TypeText(item + "\n")
            
            # Add changes as tracked changes
            selection.Font.Bold = True
            selection.TypeText("Contract Changes:\n\n")
            selection.Font.Bold = False
            
            changes = getattr(analysis, 'changes', [])
            for change in changes:
                change_type = getattr(change, 'change_type', None)
                if hasattr(change_type, 'value'):
                    change_type = change_type.value
                
                # Simulate track changes by manipulating text
                if change_type == 'deletion':
                    deleted_text = getattr(change, 'deleted_text', '')
                    if deleted_text:
                        selection.TypeText(deleted_text)
                        selection.MoveLeft(1, len(deleted_text), True)  # Select the text
                        selection.Delete()  # This creates a tracked deletion
                        
                elif change_type == 'insertion':
                    inserted_text = getattr(change, 'inserted_text', '')
                    if inserted_text:
                        selection.TypeText(inserted_text)  # This creates a tracked insertion
                        
                elif change_type == 'replacement':
                    deleted_text = getattr(change, 'deleted_text', '')
                    inserted_text = getattr(change, 'inserted_text', '')
                    
                    if deleted_text:
                        selection.TypeText(deleted_text)
                        selection.MoveLeft(1, len(deleted_text), True)
                        selection.Delete()
                    
                    if inserted_text:
                        selection.TypeText(inserted_text)
                
                selection.TypeText("\n")
            
            # Save and close
            doc.SaveAs2(os.path.abspath(output_path))
            doc.Close()
            word.Quit()
            
            return True
            
        except Exception as e:
            logger.warning(f"COM automation failed: {e}")
            return False
    
    def _add_metadata_section(self, document: Any, analysis: Any) -> None:
        """Add metadata section to document."""
        document.add_heading('Analysis Details', level=1)
        
        metadata_items = [
            f'Analysis ID: {getattr(analysis, "analysis_id", "N/A")}',
            f'Contract ID: {getattr(analysis, "contract_id", "N/A")}',
            f'Template ID: {getattr(analysis, "template_id", "N/A")}',
            f'Status: {getattr(analysis, "status", "N/A")}',
            f'Created: {getattr(analysis, "created_at", "N/A")}'
        ]
        
        for item in metadata_items:
            document.add_paragraph(item, style='List Bullet')
        
        document.add_paragraph()  # Add spacing
    
    def _add_summary_section(self, document: Any, analysis: Any) -> None:
        """Add AI analysis summary section."""
        document.add_heading('Executive Summary', level=1)
        
        llm_analysis = getattr(analysis, 'llm_analysis', None)
        if llm_analysis and 'summary' in llm_analysis:
            summary_text = llm_analysis['summary']
        else:
            summary_text = 'No AI analysis summary available for this contract analysis.'
        
        document.add_paragraph(summary_text)
        document.add_paragraph()  # Add spacing
    
    def _add_statistics_section(self, document: Any, analysis: Any) -> None:
        """Add change statistics section."""
        document.add_heading('Change Statistics', level=1)
        
        statistics = getattr(analysis, 'statistics', {})
        for key, value in statistics.items():
            display_key = key.replace('_', ' ').title()
            document.add_paragraph(f'{display_key}: {value}', style='List Bullet')
        
        document.add_paragraph()  # Add spacing
    
    def _add_changes_table(self, document: Any, analysis: Any) -> None:
        """Add detailed changes table."""
        document.add_heading('Detailed Changes', level=1)
        
        changes = getattr(analysis, 'changes', [])
        if not changes:
            document.add_paragraph('No changes detected in this analysis.')
            return
        
        # Create table
        table = document.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        header_cells = table.rows[0].cells
        headers = ['Type', 'Original', 'Modified', 'Position', 'Context']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make headers bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add change data
        formatted_changes = self.format_changes_for_display(changes)
        for change in formatted_changes:
            row_cells = table.add_row().cells
            row_cells[0].text = change['operation']
            row_cells[1].text = change['original_text'][:100] + ('...' if len(change['original_text']) > 100 else '')
            row_cells[2].text = change['modified_text'][:100] + ('...' if len(change['modified_text']) > 100 else '')
            row_cells[3].text = str(change['position'])
            row_cells[4].text = change['context'][:150] + ('...' if len(change['context']) > 150 else '')
    
    @property
    def track_changes_available(self) -> bool:
        """Check if track changes functionality is available."""
        return WIN32COM_AVAILABLE